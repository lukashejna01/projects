import os
import tkinter as tk
from tkinter import messagebox
import datetime
from docx import Document
from docx.shared import Pt

# Nastavení fontu vyplňovaného textu
def set_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = False  # Nastaví písmo na netučné

# Funkce pro vyplnění smluv
def fill_contracts():
    name = name_entry.get()
    date_of_birth = date_entry.get()
    residence = residence_entry.get()
    current_date = datetime.date.today().strftime("%d.%m.%Y")

    # Vytvoření složky pro uživatele
    user_folder = f"{name}_smlouvy"
    os.makedirs(user_folder, exist_ok=True)

    # Načtení obsahu smluv
    with open("dpp_nero.docx", 'rb') as file_dpp_template:
        contract_dpp_template = Document(file_dpp_template)
        dpp_filename = os.path.basename(file_dpp_template.name)

    with open("dpc_nero.docx", 'rb') as file_dpc_template:
        contract_dpc_template = Document(file_dpc_template)
        dpc_filename = os.path.basename(file_dpc_template.name)

    # Nahrazení proměnných ve šabloně a nastavení formátu písma
    for document, checkbox_var, filename in zip([contract_dpp_template, contract_dpc_template], [dpp_checkbox_var, dpc_checkbox_var], [dpp_filename, dpc_filename]):
        if checkbox_var.get():
            for paragraph in document.paragraphs:
                if "[jmeno]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[jmeno]", name)
                    set_font(paragraph)
                if "[narozeni]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[narozeni]", date_of_birth)
                    set_font(paragraph)
                if "[bydliste]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[bydliste]", residence)
                    set_font(paragraph)
                if "[aktualni_datum]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[aktualni_datum]", current_date)
                    set_font(paragraph)

            # Název výstupních souborů
            output_file = f"{user_folder}/{filename}_{name}.docx"
            document.save(output_file)

    messagebox.showinfo("Info", "Smlouvy byly vyplněny a uloženy.")

# Vytvoření hlavního okna
root = tk.Tk()
root.title("Vyplnění smluv")

# Vytvoření rámců pro rozdělení prostoru
input_frame = tk.Frame(root)
input_frame.pack(padx=20, pady=10, fill=tk.X)  # Zarovnání rámců na střed

checkbox_frame = tk.Frame(root)
checkbox_frame.pack(padx=20, pady=10, fill=tk.X)  # Zarovnání rámců na střed

button_frame = tk.Frame(root)
button_frame.pack(padx=20, pady=10, fill=tk.X)  # Zarovnání rámců na střed

# Vytvoření popisků a vstupních polí pro jednotlivé údaje
tk.Label(input_frame, text="Jméno:").grid(row=0, column=0, sticky="w")
name_entry = tk.Entry(input_frame, width=70)
name_entry.grid(row=0, column=1, sticky="ew")

tk.Label(input_frame, text="Datum narození:").grid(row=1, column=0, sticky="w")
date_entry = tk.Entry(input_frame, width=70)
date_entry.grid(row=1, column=1, sticky="ew")

tk.Label(input_frame, text="Bydliště:").grid(row=2, column=0, sticky="w")
residence_entry = tk.Entry(input_frame, width=70)
residence_entry.grid(row=2, column=1, sticky="ew")

# Zaškrtávací pole pro výběr smluv
tk.Label(checkbox_frame, text="Označte smlouvy, které chcete vyplnit:").grid(row=0, column=0, columnspan=2, sticky="w")
dpp_checkbox_var = tk.BooleanVar()
dpp_checkbox = tk.Checkbutton(checkbox_frame, text="DPP", variable=dpp_checkbox_var)
dpp_checkbox.grid(row=1, column=0, sticky="w")
dpc_checkbox_var = tk.BooleanVar()
dpc_checkbox = tk.Checkbutton(checkbox_frame, text="DPC", variable=dpc_checkbox_var)
dpc_checkbox.grid(row=1, column=1, sticky="w")

# Tlačítko pro vyplnění smluv
fill_button = tk.Button(button_frame, text="Vyplnit smlouvy", command=fill_contracts)
fill_button.pack()

root.mainloop()
